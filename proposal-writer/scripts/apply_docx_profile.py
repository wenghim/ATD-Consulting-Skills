#!/usr/bin/env python3
"""Apply a learned DOCX format profile to semantic Word styles."""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

try:
    from .output_path_guard import ensure_external_output
except ImportError:
    from output_path_guard import ensure_external_output


SCHEMA_VERSION = "1.0"

ROLE_TO_STYLE = {
    "normal": "Normal",
    "title": "Title",
    "subtitle": "Subtitle",
    "heading_1": "Heading 1",
    "heading_2": "Heading 2",
    "heading_3": "Heading 3",
    "quote": "Quote",
    "caption": "Caption",
    "list_bullet": "List Bullet",
    "list_number": "List Number",
    "header": "Header",
    "footer": "Footer",
}

ALIGNMENTS = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "both": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "distribute": WD_ALIGN_PARAGRAPH.DISTRIBUTE,
}


def _set_font(style: Any, values: dict[str, Any]) -> None:
    font = style.font
    if values.get("name"):
        font.name = values["name"]
    if values.get("size_pt") is not None:
        font.size = Pt(float(values["size_pt"]))
    for source_key, attribute in (
        ("bold", "bold"),
        ("italic", "italic"),
        ("underline", "underline"),
        ("caps", "all_caps"),
        ("small_caps", "small_caps"),
    ):
        if values.get(source_key) is not None:
            setattr(font, attribute, bool(values[source_key]))
    color = values.get("color")
    if isinstance(color, str) and re.fullmatch(r"[0-9A-Fa-f]{6}", color):
        font.color.rgb = RGBColor.from_string(color.upper())


def _set_paragraph_format(style: Any, values: dict[str, Any]) -> None:
    paragraph = style.paragraph_format
    if values.get("alignment") in ALIGNMENTS:
        paragraph.alignment = ALIGNMENTS[values["alignment"]]
    for source_key, attribute in (
        ("space_before_pt", "space_before"),
        ("space_after_pt", "space_after"),
        ("left_indent_pt", "left_indent"),
        ("right_indent_pt", "right_indent"),
        ("first_line_indent_pt", "first_line_indent"),
    ):
        if values.get(source_key) is not None:
            setattr(paragraph, attribute, Pt(float(values[source_key])))
    if values.get("hanging_indent_pt") is not None:
        paragraph.first_line_indent = Pt(-float(values["hanging_indent_pt"]))
    for source_key, attribute in (
        ("keep_next", "keep_with_next"),
        ("keep_lines", "keep_together"),
        ("page_break_before", "page_break_before"),
    ):
        if values.get(source_key) is not None:
            setattr(paragraph, attribute, bool(values[source_key]))


def _apply_sections(document: Document, section_profile: dict[str, Any]) -> None:
    orientation = section_profile.get("orientation")
    margins = section_profile.get("margins_in", {})
    for section in document.sections:
        if orientation == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
        elif orientation == "portrait":
            section.orientation = WD_ORIENT.PORTRAIT
        if section_profile.get("page_width_in") is not None:
            section.page_width = Inches(float(section_profile["page_width_in"]))
        if section_profile.get("page_height_in") is not None:
            section.page_height = Inches(float(section_profile["page_height_in"]))
        for source_key, attribute in (
            ("top", "top_margin"),
            ("right", "right_margin"),
            ("bottom", "bottom_margin"),
            ("left", "left_margin"),
            ("header", "header_distance"),
            ("footer", "footer_distance"),
            ("gutter", "gutter"),
        ):
            if margins.get(source_key) is not None:
                setattr(section, attribute, Inches(float(margins[source_key])))


def _apply_table_style(document: Document, profile: dict[str, Any]) -> None:
    table_profiles = profile.get("tables", [])
    if not table_profiles:
        return
    style_id = table_profiles[0].get("style_id")
    if not style_id:
        return
    style_name = profile.get("styles", {}).get(style_id, {}).get("name") or style_id
    for table in document.tables:
        try:
            table.style = style_name
        except KeyError:
            continue


def _set_cell_shading(cell: Any, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)


def _set_cell_borders(cell: Any, color: str, size: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    borders = properties.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _apply_table_design(document: Document, design: dict[str, Any]) -> None:
    if not design:
        return
    header_fill = design.get("header_fill")
    header_font_color = design.get("header_font_color")
    header_bold = design.get("header_bold")
    header_alignment = ALIGNMENTS.get(design.get("header_alignment"))
    first_column_fill = design.get("first_column_fill")
    first_column_bold = design.get("first_column_bold")
    body_fill = design.get("body_fill")
    border_color = design.get("border_color")
    border_size = int(design.get("border_size_eighth_points", 4))

    for table in document.tables:
        for row_index, row in enumerate(table.rows):
            for column_index, cell in enumerate(row.cells):
                if body_fill:
                    _set_cell_shading(cell, body_fill)
                if row_index == 0 and header_fill:
                    _set_cell_shading(cell, header_fill)
                elif column_index == 0 and first_column_fill:
                    _set_cell_shading(cell, first_column_fill)
                if border_color:
                    _set_cell_borders(cell, border_color, border_size)
                for paragraph in cell.paragraphs:
                    if row_index == 0 and header_alignment is not None:
                        paragraph.alignment = header_alignment
                    for run in paragraph.runs:
                        if row_index == 0:
                            if header_bold is not None:
                                run.bold = bool(header_bold)
                            if header_font_color:
                                run.font.color.rgb = RGBColor.from_string(header_font_color)
                        elif column_index == 0 and first_column_bold is not None:
                            run.bold = bool(first_column_bold)


def apply_profile(
    source: Path | str,
    profile_path: Path | str,
    output: Path | str,
) -> Path:
    source = Path(source)
    profile_path = Path(profile_path)
    output = ensure_external_output(output)
    if source.resolve() == output.resolve():
        raise ValueError("Profile application requires a distinct output path; in-place edits are forbidden")
    if source.suffix.lower() != ".docx" or not source.is_file():
        raise ValueError(f"Source is not a readable DOCX: {source}")
    profile = json.loads(profile_path.read_text(encoding="utf-8"))
    if profile.get("schema_version") != SCHEMA_VERSION:
        raise ValueError(f"Unsupported profile schema: {profile.get('schema_version')}")

    document = Document(source)
    roles = profile.get("semantic_roles", {})
    for role, target_style_name in ROLE_TO_STYLE.items():
        values = roles.get(role)
        if not values:
            continue
        try:
            style = document.styles[target_style_name]
        except KeyError:
            continue
        _set_font(style, values.get("font", {}))
        _set_paragraph_format(style, values.get("paragraph", {}))

    sections = profile.get("sections", [])
    if sections:
        _apply_sections(document, sections[0])
    _apply_table_style(document, profile)
    _apply_table_design(document, profile.get("table_design", {}))

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return output


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("source", type=Path)
    parser.add_argument("profile", type=Path)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    print(apply_profile(args.source, args.profile, args.output))


if __name__ == "__main__":
    main()
