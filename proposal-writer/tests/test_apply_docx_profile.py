import json
import tempfile
import unittest
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document

from scripts.analyze_docx_format import analyze_docx
from tests.docx_fixtures import build_format_fixture, logical_text


class ApplyDocxProfileTests(unittest.TestCase):
    def test_applies_semantic_styles_and_preserves_text(self):
        from scripts.apply_docx_profile import apply_profile

        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            reference = build_format_fixture(root / "reference.docx")
            profile_path = root / "profile.json"
            profile_path.write_text(
                json.dumps(analyze_docx(reference, "Fixture Blue")), encoding="utf-8"
            )

            target = Document()
            target.add_paragraph("Target title", style="Title")
            target.add_heading("Target heading", level=1)
            target.add_paragraph("Target body text")
            target.add_paragraph("Target list item", style="List Number")
            table = target.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "A"
            table.cell(0, 1).text = "B"
            table.cell(1, 0).text = "C"
            table.cell(1, 1).text = "D"
            target.sections[0].header.paragraphs[0].text = "Target Header"
            target.sections[0].footer.paragraphs[0].text = "Target Footer"
            target_path = root / "target.docx"
            target.save(target_path)
            before = logical_text(target_path)

            output_path = apply_profile(
                target_path, profile_path, root / "styled-target.docx"
            )
            after = logical_text(output_path)
            output = Document(output_path)

        self.assertEqual(before, after)
        self.assertEqual(output.styles["Title"].font.size.pt, 24.0)
        self.assertEqual(output.styles["Heading 1"].font.size.pt, 16.0)
        self.assertEqual(output.styles["Normal"].font.size.pt, 10.5)
        self.assertAlmostEqual(output.sections[0].top_margin.inches, 0.8, places=1)
        self.assertNotEqual(output_path, target_path)

    def test_rejects_in_place_output(self):
        from scripts.apply_docx_profile import apply_profile

        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            source = build_format_fixture(root / "source.docx")
            profile = root / "profile.json"
            profile.write_text(
                json.dumps(analyze_docx(source, "Fixture")), encoding="utf-8"
            )
            with self.assertRaisesRegex(ValueError, "distinct output"):
                apply_profile(source, profile, source)

    def test_rejects_unsupported_schema(self):
        from scripts.apply_docx_profile import apply_profile

        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            source = build_format_fixture(root / "source.docx")
            profile = root / "profile.json"
            profile.write_text('{"schema_version":"9.9"}', encoding="utf-8")
            with self.assertRaisesRegex(ValueError, "schema"):
                apply_profile(source, profile, root / "output.docx")

    def test_applies_direct_table_design_tokens(self):
        from scripts.apply_docx_profile import apply_profile

        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            target = Document()
            table = target.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Header A"
            table.cell(0, 1).text = "Header B"
            table.cell(1, 0).text = "Label"
            table.cell(1, 1).text = "Value"
            source = root / "source.docx"
            target.save(source)
            profile = root / "profile.json"
            profile.write_text(
                json.dumps(
                    {
                        "schema_version": "1.0",
                        "semantic_roles": {},
                        "sections": [],
                        "tables": [],
                        "table_design": {
                            "header_fill": "C00000",
                            "header_font_color": "FFFFFF",
                            "header_bold": True,
                            "first_column_fill": "FAF0F0",
                            "border_color": "000000",
                            "border_size_eighth_points": 4,
                        },
                    }
                ),
                encoding="utf-8",
            )

            output = apply_profile(source, profile, root / "output.docx")
            with zipfile.ZipFile(output) as package:
                root_xml = ET.fromstring(package.read("word/document.xml"))

        namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        cells = root_xml.findall(".//w:tbl/w:tr/w:tc", namespace)
        header_fill = cells[0].find("w:tcPr/w:shd", namespace)
        first_column_fill = cells[2].find("w:tcPr/w:shd", namespace)
        header_color = cells[0].find(".//w:rPr/w:color", namespace)
        self.assertIn("C00000", header_fill.attrib.values())
        self.assertIn("FAF0F0", first_column_fill.attrib.values())
        self.assertIn("FFFFFF", header_color.attrib.values())


if __name__ == "__main__":
    unittest.main()
