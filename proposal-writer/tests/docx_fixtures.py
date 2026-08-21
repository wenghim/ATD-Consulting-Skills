from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


def build_format_fixture(path: Path) -> Path:
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)

    title_style = document.styles["Title"]
    title_style.font.name = "Aptos Display"
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(31, 78, 121)

    heading = document.styles["Heading 1"]
    heading.font.name = "Aptos Display"
    heading.font.size = Pt(16)
    heading.font.bold = True
    heading.font.color.rgb = RGBColor(47, 84, 150)
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)

    document.add_paragraph("Fixture Document", style="Title")
    document.add_paragraph("Format Analysis", style="Subtitle")
    document.add_heading("Overview", level=1)

    paragraph = document.add_paragraph("Standard body text with ")
    exception = paragraph.add_run("direct emphasis")
    exception.bold = True
    exception.font.color.rgb = RGBColor(192, 0, 0)
    paragraph.add_run(" preserved.")
    document.add_paragraph("First numbered item", style="List Number")

    table = document.add_table(rows=2, cols=2)
    table.style = "Light Shading Accent 1"
    table.cell(0, 0).text = "Label"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Theme"
    table.cell(1, 1).text = "Blue"

    header = section.header.paragraphs[0]
    header.text = "MS Master Header"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer = section.footer.paragraphs[0]
    footer.text = "MS Master Footer"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    image_path = path.with_suffix(".png")
    Image.new("RGB", (20, 12), color=(31, 78, 121)).save(image_path)
    picture_paragraph = document.add_paragraph()
    run = picture_paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(0.5))
    doc_pr = picture_paragraph._p.xpath(".//wp:docPr")[0]
    doc_pr.set("name", "Brand sample")
    doc_pr.set("descr", "Blue rectangular brand sample")
    document.add_paragraph("Figure 1. Brand color sample", style="Caption")

    document.save(path)
    return path


def logical_text(path: Path) -> list[str]:
    document = Document(path)
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            values.extend(cell.text for cell in row.cells)
    for section in document.sections:
        values.extend(paragraph.text for paragraph in section.header.paragraphs)
        values.extend(paragraph.text for paragraph in section.footer.paragraphs)
    return values

