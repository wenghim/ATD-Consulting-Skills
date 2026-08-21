import json
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document


NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def highlighted_text(path: Path) -> list[str]:
    from xml.etree import ElementTree as ET

    values = []
    with zipfile.ZipFile(path) as package:
        for part in ("word/document.xml", "word/header1.xml", "word/footer1.xml"):
            if part not in package.namelist():
                continue
            root = ET.fromstring(package.read(part))
            for run in root.findall(".//w:r", NS):
                if run.find("w:rPr/w:highlight", NS) is None:
                    continue
                text = "".join(node.text or "" for node in run.findall(".//w:t", NS))
                if text:
                    values.append(text)
    return values


class FlagDocxVariablesTests(unittest.TestCase):
    def test_highlights_markers_and_manifest_terms_without_changing_text(self):
        from scripts.flag_docx_variables import flag_docx_variables

        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            document = Document()
            document.add_paragraph("Prepared for PUB")
            document.add_paragraph("Version: Ver 1.0")
            table = document.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "Contact: {{CONTACT_PERSON}}"
            document.sections[0].header.paragraphs[0].text = "Year {{DOCUMENT_YEAR}}"
            document.sections[0].footer.paragraphs[0].text = "Page footer"
            source = root / "source.docx"
            document.save(source)
            manifest = root / "manifest.json"
            manifest.write_text(
                json.dumps(
                    {
                        "highlight_color": "FFFF00",
                        "marker_regex": r"\{\{[A-Z][A-Z0-9_]*\}\}",
                        "observed_terms": ["PUB"],
                        "candidate_regexes": [r"\bVer\s+\d+(?:\.\d+)+\b"],
                    }
                ),
                encoding="utf-8",
            )

            output = flag_docx_variables(source, manifest, root / "flagged.docx")
            before = Document(source)
            after = Document(output)
            marked = highlighted_text(output)

        self.assertEqual(before.paragraphs[0].text, after.paragraphs[0].text)
        self.assertIn("PUB", marked)
        self.assertIn("Ver 1.0", marked)
        self.assertNotIn("Prepared for PUB", marked)
        self.assertNotIn("Version: Ver 1.0", marked)
        self.assertTrue(any("{{CONTACT_PERSON}}" in value for value in marked))
        self.assertTrue(any("{{DOCUMENT_YEAR}}" in value for value in marked))

    def test_rejects_in_place_output(self):
        from scripts.flag_docx_variables import flag_docx_variables

        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            document = Document()
            document.add_paragraph("{{CLIENT_NAME}}")
            source = root / "source.docx"
            document.save(source)
            manifest = root / "manifest.json"
            manifest.write_text(
                '{"highlight_color":"FFFF00","marker_regex":"x","observed_terms":[]}',
                encoding="utf-8",
            )
            with self.assertRaisesRegex(ValueError, "distinct output"):
                flag_docx_variables(source, manifest, source)


if __name__ == "__main__":
    unittest.main()
